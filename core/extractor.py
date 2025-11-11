"""
Metadata Extractor Module
Extracts metadata from various file types including images, documents, and media files.
"""

import os
import mimetypes
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional
import hashlib

# Image processing
try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    import exifread
    IMAGE_SUPPORT = True
except ImportError:
    IMAGE_SUPPORT = False

# Document processing
try:
    import PyPDF2
    import fitz  # PyMuPDF
    from docx import Document as DocxDocument
    import openpyxl
    from pptx import Presentation
    DOCUMENT_SUPPORT = True
except ImportError:
    DOCUMENT_SUPPORT = False

# Media processing
try:
    from mutagen import File as MutagenFile
    import subprocess
    import json
    MEDIA_SUPPORT = True
except ImportError:
    MEDIA_SUPPORT = False


class MetadataExtractor:
    """
    Main class for extracting metadata from various file types.
    Supports images, documents (PDF, DOCX, XLSX, PPTX), and media files.
    """
    
    def __init__(self, file_path: str):
        """
        Initialize the MetadataExtractor with a file path.
        
        Args:
            file_path (str): Path to the file to analyze
        """
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        self.mime_type, _ = mimetypes.guess_type(str(self.file_path))
        self.file_type = self._determine_file_type()
    
    def _determine_file_type(self) -> str:
        """Determine the category of the file."""
        ext = self.file_path.suffix.lower()
        
        image_exts = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif', '.heic', '.heif'}
        document_exts = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx'}
        video_exts = {'.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv', '.webm'}
        audio_exts = {'.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a', '.wma'}
        
        if ext in image_exts:
            return 'image'
        elif ext in document_exts:
            return 'document'
        elif ext in video_exts:
            return 'video'
        elif ext in audio_exts:
            return 'audio'
        else:
            return 'unknown'
    
    def extract_all(self) -> Dict[str, Any]:
        """
        Extract all available metadata from the file.
        
        Returns:
            Dict containing all extracted metadata
        """
        metadata = {
            'file_info': self._extract_file_system_metadata(),
            'file_type': self.file_type,
            'mime_type': self.mime_type,
            'extraction_time': datetime.now().isoformat()
        }
        
        # Extract type-specific metadata
        if self.file_type == 'image' and IMAGE_SUPPORT:
            metadata['image_metadata'] = self._extract_image_metadata()
            metadata['exif_data'] = self._extract_exif_data()
            metadata['gps_data'] = self._extract_gps_data()
        elif self.file_type == 'document' and DOCUMENT_SUPPORT:
            metadata['document_metadata'] = self._extract_document_metadata()
        elif self.file_type in ['video', 'audio'] and MEDIA_SUPPORT:
            metadata['media_metadata'] = self._extract_media_metadata()
        
        return metadata
    
    def _extract_file_system_metadata(self) -> Dict[str, Any]:
        """Extract file system metadata."""
        stat = self.file_path.stat()
        
        return {
            'filename': self.file_path.name,
            'full_path': str(self.file_path.absolute()),
            'extension': self.file_path.suffix,
            'size_bytes': stat.st_size,
            'size_human': self._format_size(stat.st_size),
            'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'accessed': datetime.fromtimestamp(stat.st_atime).isoformat(),
            'md5_hash': self._calculate_hash('md5'),
            'sha256_hash': self._calculate_hash('sha256'),
            'note': 'File system dates may differ from original dates if file was copied/moved'
        }
    
    def _calculate_hash(self, algorithm: str = 'md5') -> str:
        """Calculate file hash."""
        hash_func = hashlib.new(algorithm)
        try:
            with open(self.file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b''):
                    hash_func.update(chunk)
            return hash_func.hexdigest()
        except Exception as e:
            return f"Error: {str(e)}"
    
    def _format_size(self, size_bytes: int) -> str:
        """Format file size in human-readable format."""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} PB"
    
    def _extract_image_metadata(self) -> Dict[str, Any]:
        """Extract basic image metadata using Pillow."""
        if not IMAGE_SUPPORT:
            return {'error': 'Image support not available'}
        
        try:
            with Image.open(self.file_path) as img:
                return {
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'width': img.width,
                    'height': img.height,
                    'info': dict(img.info) if hasattr(img, 'info') else {}
                }
        except Exception as e:
            return {'error': str(e)}
    
    def _extract_exif_data(self) -> Dict[str, Any]:
        """Extract EXIF data from images."""
        if not IMAGE_SUPPORT:
            return {'error': 'Image support not available'}
        
        exif_data = {}
        
        try:
            # Try with Pillow first
            with Image.open(self.file_path) as img:
                exif = img.getexif()
                if exif:
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        exif_data[tag] = str(value)
            
            # Also try with exifread for more detailed data
            with open(self.file_path, 'rb') as f:
                tags = exifread.process_file(f, details=False)
                for tag, value in tags.items():
                    if tag not in ['JPEGThumbnail', 'TIFFThumbnail']:
                        exif_data[tag] = str(value)
            
            # Extract and parse original dates if present
            if exif_data:
                original_dates = self._parse_exif_dates(exif_data)
                if original_dates:
                    exif_data['parsed_dates'] = original_dates
            
            return exif_data if exif_data else {'info': 'No EXIF data found'}
        
        except Exception as e:
            return {'error': str(e)}
    
    def _parse_exif_dates(self, exif_data: Dict[str, Any]) -> Dict[str, str]:
        """Parse and extract date fields from EXIF data."""
        dates = {}
        
        # Common EXIF date fields
        date_fields = {
            'DateTime': 'File Change Date',
            'DateTimeOriginal': 'Original Date (when photo was taken)',
            'DateTimeDigitized': 'Digitized Date (when photo was saved)',
            'EXIF DateTimeOriginal': 'Original Date (when photo was taken)',
            'EXIF DateTimeDigitized': 'Digitized Date (when photo was saved)',
            'Image DateTime': 'Image Date',
        }
        
        for exif_key, description in date_fields.items():
            if exif_key in exif_data:
                try:
                    date_str = str(exif_data[exif_key])
                    # Try to parse EXIF date format: YYYY:MM:DD HH:MM:SS
                    if ':' in date_str:
                        # Convert EXIF format to ISO format
                        date_str_clean = date_str.replace(':', '-', 2)  # Replace first two colons
                        dates[description] = date_str_clean
                    else:
                        dates[description] = date_str
                except Exception:
                    pass
        
        return dates
    
    def _extract_gps_data(self) -> Optional[Dict[str, Any]]:
        """Extract GPS coordinates from image EXIF data."""
        if not IMAGE_SUPPORT:
            return {'error': 'Image support not available'}
        
        try:
            with Image.open(self.file_path) as img:
                exif = img.getexif()
                if not exif:
                    return None
                
                gps_info = exif.get_ifd(0x8825)  # GPS IFD
                if not gps_info:
                    return None
                
                gps_data = {}
                for tag_id, value in gps_info.items():
                    tag = GPSTAGS.get(tag_id, tag_id)
                    gps_data[tag] = value
                
                # Convert to decimal degrees if coordinates are present
                if 'GPSLatitude' in gps_data and 'GPSLongitude' in gps_data:
                    lat = self._convert_to_degrees(gps_data['GPSLatitude'])
                    lon = self._convert_to_degrees(gps_data['GPSLongitude'])
                    
                    # Apply direction
                    if gps_data.get('GPSLatitudeRef') == 'S':
                        lat = -lat
                    if gps_data.get('GPSLongitudeRef') == 'W':
                        lon = -lon
                    
                    gps_data['latitude_decimal'] = lat
                    gps_data['longitude_decimal'] = lon
                    gps_data['coordinates'] = f"{lat}, {lon}"
                
                # Extract altitude if present
                if 'GPSAltitude' in gps_data:
                    altitude = float(gps_data['GPSAltitude'])
                    gps_data['altitude_meters'] = altitude
                
                return gps_data
        
        except Exception as e:
            return {'error': str(e)}
    
    def _convert_to_degrees(self, value) -> float:
        """Convert GPS coordinates to decimal degrees."""
        d, m, s = value
        return float(d) + float(m) / 60.0 + float(s) / 3600.0
    
    def _extract_document_metadata(self) -> Dict[str, Any]:
        """Extract metadata from documents (PDF, DOCX, XLSX, PPTX)."""
        if not DOCUMENT_SUPPORT:
            return {'error': 'Document support not available'}
        
        ext = self.file_path.suffix.lower()
        
        try:
            if ext == '.pdf':
                return self._extract_pdf_metadata()
            elif ext in ['.doc', '.docx']:
                return self._extract_docx_metadata()
            elif ext in ['.xls', '.xlsx']:
                return self._extract_xlsx_metadata()
            elif ext in ['.ppt', '.pptx']:
                return self._extract_pptx_metadata()
            else:
                return {'error': 'Unsupported document type'}
        except Exception as e:
            return {'error': str(e)}
    
    def _extract_pdf_metadata(self) -> Dict[str, Any]:
        """Extract metadata from PDF files."""
        metadata = {}
        
        try:
            # Try with PyMuPDF (fitz) - more reliable
            doc = fitz.open(self.file_path)
            metadata['page_count'] = doc.page_count
            metadata['metadata'] = doc.metadata
            
            # Extract text from first page as sample
            if doc.page_count > 0:
                first_page = doc[0]
                metadata['first_page_text_sample'] = first_page.get_text()[:500]
            
            doc.close()
        except:
            # Fallback to PyPDF2
            try:
                with open(self.file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    metadata['page_count'] = len(pdf_reader.pages)
                    
                    if pdf_reader.metadata:
                        metadata['metadata'] = {
                            key: str(value) for key, value in pdf_reader.metadata.items()
                        }
            except Exception as e:
                metadata['error'] = str(e)
        
        return metadata
    
    def _extract_docx_metadata(self) -> Dict[str, Any]:
        """Extract metadata from DOCX files."""
        try:
            doc = DocxDocument(self.file_path)
            core_props = doc.core_properties
            
            return {
                'author': core_props.author,
                'title': core_props.title,
                'subject': core_props.subject,
                'keywords': core_props.keywords,
                'comments': core_props.comments,
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'last_modified_by': core_props.last_modified_by,
                'revision': core_props.revision,
                'category': core_props.category,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
            }
        except Exception as e:
            return {'error': str(e)}
    
    def _extract_xlsx_metadata(self) -> Dict[str, Any]:
        """Extract metadata from XLSX files."""
        try:
            wb = openpyxl.load_workbook(self.file_path)
            props = wb.properties
            
            return {
                'creator': props.creator,
                'title': props.title,
                'subject': props.subject,
                'description': props.description,
                'keywords': props.keywords,
                'category': props.category,
                'created': props.created.isoformat() if props.created else None,
                'modified': props.modified.isoformat() if props.modified else None,
                'last_modified_by': props.lastModifiedBy,
                'sheet_count': len(wb.sheetnames),
                'sheet_names': wb.sheetnames,
            }
        except Exception as e:
            return {'error': str(e)}
    
    def _extract_pptx_metadata(self) -> Dict[str, Any]:
        """Extract metadata from PPTX files."""
        try:
            prs = Presentation(self.file_path)
            core_props = prs.core_properties
            
            return {
                'author': core_props.author,
                'title': core_props.title,
                'subject': core_props.subject,
                'keywords': core_props.keywords,
                'comments': core_props.comments,
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'last_modified_by': core_props.last_modified_by,
                'revision': core_props.revision,
                'category': core_props.category,
                'slide_count': len(prs.slides),
            }
        except Exception as e:
            return {'error': str(e)}
    
    def _extract_media_metadata(self) -> Dict[str, Any]:
        """Extract metadata from media files (audio/video)."""
        if not MEDIA_SUPPORT:
            return {'error': 'Media support not available'}
        
        metadata = {}
        
        # Try mutagen first for audio
        try:
            audio = MutagenFile(self.file_path)
            if audio:
                metadata['mutagen_data'] = {
                    'length': audio.info.length if hasattr(audio.info, 'length') else None,
                    'bitrate': audio.info.bitrate if hasattr(audio.info, 'bitrate') else None,
                    'sample_rate': audio.info.sample_rate if hasattr(audio.info, 'sample_rate') else None,
                    'channels': audio.info.channels if hasattr(audio.info, 'channels') else None,
                }
                
                # Extract tags
                if audio.tags:
                    metadata['tags'] = {str(k): str(v) for k, v in audio.tags.items()}
        except Exception as e:
            metadata['mutagen_error'] = str(e)
        
        # Try ffprobe for detailed video/audio info
        try:
            ffprobe_data = self._extract_ffprobe_metadata()
            if ffprobe_data:
                metadata['ffprobe_data'] = ffprobe_data
        except Exception as e:
            metadata['ffprobe_error'] = str(e)
        
        return metadata
    
    def _extract_ffprobe_metadata(self) -> Optional[Dict[str, Any]]:
        """Extract metadata using ffprobe (part of ffmpeg)."""
        try:
            cmd = [
                'ffprobe',
                '-v', 'quiet',
                '-print_format', 'json',
                '-show_format',
                '-show_streams',
                str(self.file_path)
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0:
                return json.loads(result.stdout)
            else:
                return None
        except FileNotFoundError:
            return {'error': 'ffprobe not found. Please install ffmpeg.'}
        except Exception as e:
            return {'error': str(e)}
