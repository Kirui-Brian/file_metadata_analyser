#!/usr/bin/env python3
"""
Sample File Creator
Creates basic sample files for testing the File Metadata Analyser.
"""

import os
from pathlib import Path
from datetime import datetime

def create_sample_image():
    """Create a simple test image."""
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # Create a simple image
        img = Image.new('RGB', (800, 600), color='lightblue')
        draw = ImageDraw.Draw(img)
        
        # Add text
        text = f"Sample Test Image\nCreated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        draw.text((50, 50), text, fill='darkblue')
        
        # Add some shapes
        draw.rectangle([100, 200, 700, 500], outline='red', width=3)
        draw.ellipse([200, 250, 600, 450], outline='green', width=3)
        
        # Save
        output_path = Path("samples/test_image.jpg")
        img.save(output_path, 'JPEG', quality=95)
        
        print(f"✓ Created: {output_path}")
        return True
    except ImportError:
        print("✗ PIL/Pillow not available. Cannot create sample image.")
        return False
    except Exception as e:
        print(f"✗ Error creating image: {e}")
        return False


def create_sample_pdf():
    """Create a simple test PDF."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        output_path = Path("samples/test_document.pdf")
        c = canvas.Canvas(str(output_path), pagesize=letter)
        
        # Add content
        c.setTitle("Sample Test Document")
        c.setAuthor("File Metadata Analyser Test Suite")
        c.setSubject("Digital Forensics Testing")
        
        c.drawString(100, 750, "Sample Test Document")
        c.drawString(100, 730, "=" * 50)
        c.drawString(100, 700, f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        c.drawString(100, 680, "")
        c.drawString(100, 660, "This is a test PDF document created for demonstrating")
        c.drawString(100, 640, "the File Metadata Analyser tool.")
        c.drawString(100, 620, "")
        c.drawString(100, 600, "Metadata Information:")
        c.drawString(100, 580, "- Author: File Metadata Analyser Test Suite")
        c.drawString(100, 560, "- Subject: Digital Forensics Testing")
        c.drawString(100, 540, "- Creation Date: Embedded in file")
        
        c.save()
        
        print(f"✓ Created: {output_path}")
        return True
    except ImportError:
        print("✗ reportlab not available. Cannot create sample PDF.")
        print("  Install with: pip install reportlab")
        return False
    except Exception as e:
        print(f"✗ Error creating PDF: {e}")
        return False


def create_sample_docx():
    """Create a simple test Word document."""
    try:
        from docx import Document
        
        doc = Document()
        
        # Set metadata
        core_props = doc.core_properties
        core_props.author = "File Metadata Analyser Test Suite"
        core_props.title = "Sample Test Document"
        core_props.subject = "Digital Forensics Testing"
        core_props.keywords = "metadata, forensics, testing"
        core_props.comments = "This is a test document for demonstration purposes"
        core_props.category = "Testing"
        
        # Add content
        doc.add_heading('Sample Test Document', 0)
        doc.add_paragraph(f'Created: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        doc.add_heading('Purpose', 1)
        doc.add_paragraph(
            'This is a test Word document created for demonstrating the '
            'File Metadata Analyser tool. It contains various metadata fields '
            'including author, title, subject, keywords, and comments.'
        )
        doc.add_paragraph('')
        doc.add_heading('Metadata Contents', 1)
        doc.add_paragraph('• Author: File Metadata Analyser Test Suite')
        doc.add_paragraph('• Title: Sample Test Document')
        doc.add_paragraph('• Subject: Digital Forensics Testing')
        doc.add_paragraph('• Keywords: metadata, forensics, testing')
        doc.add_paragraph('• Category: Testing')
        
        output_path = Path("samples/test_document.docx")
        doc.save(output_path)
        
        print(f"✓ Created: {output_path}")
        return True
    except ImportError:
        print("✗ python-docx not available. Cannot create sample DOCX.")
        return False
    except Exception as e:
        print(f"✗ Error creating DOCX: {e}")
        return False


def create_sample_text():
    """Create a simple text file."""
    try:
        output_path = Path("samples/test_file.txt")
        
        content = f"""Sample Test Text File
{'='*50}
Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

This is a simple text file for testing the File Metadata Analyser.

Text files have minimal embedded metadata, but still contain:
- File system metadata (creation time, modification time, access time)
- File size
- File permissions

This file can be used to demonstrate basic file system metadata extraction.
"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"✓ Created: {output_path}")
        return True
    except Exception as e:
        print(f"✗ Error creating text file: {e}")
        return False


def create_readme():
    """Create README with instructions."""
    content = """# Test Sample Files

This directory contains automatically generated sample files for testing the File Metadata Analyser.

## Files Created

- `test_image.jpg` - Simple test image (no GPS data)
- `test_document.pdf` - PDF with author metadata
- `test_document.docx` - Word document with metadata
- `test_file.txt` - Simple text file

## Important Notes

These files contain:
✓ Safe, non-sensitive test data
✓ Embedded metadata for testing extraction
✓ Various file formats for comprehensive testing

## Adding GPS Data

To test GPS extraction, you need to:
1. Take a photo with a smartphone (GPS-enabled)
2. Or use an EXIF editor to add GPS coordinates to test_image.jpg

## Testing

Run the Analyser:
```bash
python metadata_Analyser.py --file samples/test_image.jpg --analyze
python metadata_Analyser.py --file samples/test_document.pdf --report json
python metadata_Analyser.py --directory samples/ --analyze
```

Run the demo:
```bash
python demo.py
```
"""
    
    try:
        output_path = Path("samples/GENERATED_SAMPLES_README.txt")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"✓ Created: {output_path}")
        return True
    except Exception as e:
        print(f"✗ Error creating README: {e}")
        return False


def main():
    """Main execution."""
    print("=" * 60)
    print("Sample File Creator for File Metadata Analyser")
    print("=" * 60)
    print()
    
    # Ensure samples directory exists
    samples_dir = Path("samples")
    samples_dir.mkdir(exist_ok=True)
    
    print("Creating sample files...\n")
    
    # Create samples
    results = []
    results.append(create_sample_image())
    results.append(create_sample_pdf())
    results.append(create_sample_docx())
    results.append(create_sample_text())
    results.append(create_readme())
    
    print()
    print("=" * 60)
    print(f"Summary: {sum(results)}/{len(results)} files created successfully")
    print("=" * 60)
    print()
    
    if sum(results) < len(results):
        print("⚠ Some files could not be created due to missing dependencies.")
        print("  Install all dependencies: pip install -r requirements.txt")
        print()
    
    print("Next steps:")
    print("1. Review the created files in the samples/ directory")
    print("2. Add a photo with GPS data from your smartphone (optional)")
    print("3. Run: python metadata_Analyser.py --file samples/test_image.jpg")
    print("4. Or run the full demo: python demo.py")
    print()


if __name__ == '__main__':
    main()
